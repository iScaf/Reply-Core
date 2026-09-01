# -*- coding: utf-8 -*-
"""document_chunking 纯函数测试：解析/切块/面包屑/父子结构（不依赖数据库）。"""
import asyncio
import io

import pytest

from src.chat.features.tutorial_search.services.document_chunking import (
    CHUNK_SIZE,
    MIN_CHUNK_CHARS,
    build_blocks,
    chunk_document,
    chunk_markdown,
    parse_to_markdown,
    split_sections,
)


LONG_SECTION_BODY = (
    "这是一段用于撑长篇幅的正文内容，用来确保该章节超过切块阈值。"
    "检索系统会把长章节拆分为父块与若干子块，子块参与向量与全文检索。"
) * 24  # ≈ 1.6k 字符 > CHUNK_SIZE，可拆出多个子块


def _long_md() -> str:
    return (
        "# 部署手册\n\n"
        "## 第二章 环境准备\n\n"
        f"{LONG_SECTION_BODY}\n\n"
        "### 2.1 安装依赖\n\n"
        "短章节内容，不足一个块。\n\n"
        f"{LONG_SECTION_BODY}\n\n"
        "## 第三章 回滚\n\n"
        f"| 列A | 列B |\n| --- | --- |\n"
        + "\n".join(f"| a{i} | b{i} |" for i in range(30))
        + "\n"
    )


# --- MD 切块：面包屑与父子结构 ---


def test_split_sections_breadcrumb():
    sections = split_sections(_long_md())
    paths = {s.path for s in sections}
    assert "部署手册 > 第二章 环境准备" in paths
    assert "部署手册 > 第二章 环境准备 > 2.1 安装依赖" in paths
    assert "部署手册 > 第三章 回滚" in paths


def test_build_blocks_parent_child_structure():
    blocks = build_blocks(split_sections(_long_md()))
    parents = [b for b in blocks if b.parent_text]
    singles = [b for b in blocks if not b.parent_text]
    # 长节应有父子结构，且子块均低于阈值上限
    assert parents, "长节应产出父块"
    for p in parents:
        assert len(p.children) >= 2
        assert all(len(c) <= CHUNK_SIZE + CHUNK_SIZE // 2 for c in p.children)
        assert all(len(c) >= MIN_CHUNK_CHARS for c in p.children)
    # 短节/表格节为单块
    assert singles, "短节与表格节应为单块"


def test_build_blocks_table_section_not_split():
    blocks = build_blocks(split_sections(_long_md()))
    table_blocks = [b for b in blocks if "列A" in (b.parent_text or b.children[0])]
    assert len(table_blocks) == 1
    assert table_blocks[0].parent_text is None  # 表格节整块，不拆父子


def test_short_markdown_single_block():
    """纯文本短文档（如 Discord 提交的教程）→ 单块，无父子（与旧行为一致）。"""
    result = chunk_markdown("这是一篇很短的教程内容，直接整篇作为一个块即可。")
    assert len(result.blocks) == 1
    assert result.blocks[0].parent_text is None
    assert result.blocks[0].path == ""
    assert result.blocks[0].children == ["这是一篇很短的教程内容，直接整篇作为一个块即可。"]


def test_chunk_markdown_empty_raises():
    with pytest.raises(ValueError):
        chunk_markdown("# 只有标题\n\n短")


def test_noise_section_filtered():
    md = "# 标题\n\n## 空节\n\nx\n\n## 正常节\n\n" + "正常内容" * 100
    blocks = build_blocks(split_sections(md))
    assert all("空节" not in b.path for b in blocks)


# --- 文件解析：docx / xlsx / pdf / 拒绝 ---


def _make_docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("部署手册", level=1)
    document.add_heading("第二章 环境准备", level=2)
    document.add_paragraph(LONG_SECTION_BODY)
    document.add_paragraph(LONG_SECTION_BODY)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "版本"
    table.cell(1, 0).text = "python"
    table.cell(1, 1).text = "3.14"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes() -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "版本清单"
    sheet.append(["名称", "版本"])
    sheet.append(["python", "3.14"])
    sheet.append(["postgresql", "16"])
    buf = io.BytesIO()
    workbook.save(buf)
    return buf.getvalue()


def _make_pdf_bytes() -> bytes:
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Reply-Core upload test\nSecond line of tutorial.")
    buf = io.BytesIO(doc.tobytes())
    doc.close()
    return buf.getvalue()


def test_parse_docx():
    md = asyncio.run(parse_to_markdown("部署手册.docx", _make_docx_bytes()))
    assert "# 部署手册" in md
    assert "## 第二章 环境准备" in md
    assert "| 名称 | 版本 |" in md
    sections = split_sections(md)
    assert any(s.path == "部署手册 > 第二章 环境准备" for s in sections)


def test_parse_xlsx():
    md = asyncio.run(parse_to_markdown("清单.xlsx", _make_xlsx_bytes()))
    assert "## Sheet: 版本清单" in md
    assert "| postgresql | 16 |" in md


def test_parse_pdf_sync_api():
    """parse_to_markdown 为 async：文本型 PDF 走 pymupdf4llm 快速路径。"""
    import asyncio

    md = asyncio.run(parse_to_markdown("测试.pdf", _make_pdf_bytes()))
    assert "Reply-Core upload test" in md


@pytest.mark.asyncio
async def test_parse_pdf_scan_page_routes_to_vision():
    """扫描页（文字层稀薄且含图片）应路由到 Vision 解析（注入 fake parser 验证）。"""
    import pymupdf

    # 构造"扫描页"：无文字层、整页一张图
    doc = pymupdf.open()
    page = doc.new_page(width=400, height=300)
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 380, 280))
    pix.clear_with(200)  # 灰色底
    page.insert_image(pymupdf.Rect(10, 10, 390, 290), pixmap=pix)
    buf = io.BytesIO(doc.tobytes())
    doc.close()
    raw = buf.getvalue()

    captured = {}

    async def fake_vision(png_bytes: bytes, page_no: int):
        captured["page_no"] = page_no
        captured["png_size"] = len(png_bytes)
        return "扫描页的视觉转录内容"

    md = ""
    # 通过注入 fake parser 测试路由，不真调 ai_service
    from src.chat.features.tutorial_search.services.document_chunking import (
        _pdf_to_markdown,
    )

    md = await _pdf_to_markdown(raw, vision_parser=fake_vision)
    assert "扫描页的视觉转录内容" in md
    assert captured["page_no"] == 1
    assert captured["png_size"] > 0


@pytest.mark.asyncio
async def test_parse_pdf_mixed_pages():
    """混合文档：正常页走 pymupdf4llm，扫描页走 Vision，按页序拼接。"""
    import pymupdf

    doc = pymupdf.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "normal text page one")
    p2 = doc.new_page(width=400, height=300)
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 380, 280))
    pix.clear_with(200)
    p2.insert_image(pymupdf.Rect(10, 10, 390, 290), pixmap=pix)
    buf = io.BytesIO(doc.tobytes())
    doc.close()

    async def fake_vision(png_bytes: bytes, page_no: int):
        return f"视觉解析第{page_no}页"

    from src.chat.features.tutorial_search.services.document_chunking import (
        _pdf_to_markdown,
    )

    md = await _pdf_to_markdown(buf.getvalue(), vision_parser=fake_vision)
    assert "normal text page one" in md  # 正常页走快速路径
    assert "视觉解析第2页" in md  # 扫描页走 Vision
    # 页序保持：正常页内容在扫描页之前
    assert md.index("normal text") < md.index("视觉解析第2页")


@pytest.mark.asyncio
async def test_scan_page_skipped_when_vision_fails():
    """Vision 解析失败时该页标注跳过，不阻断整本解析。"""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page(width=400, height=300)
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 380, 280))
    pix.clear_with(200)
    page.insert_image(pymupdf.Rect(10, 10, 390, 290), pixmap=pix)
    buf = io.BytesIO(doc.tobytes())
    doc.close()

    async def failing_vision(png_bytes: bytes, page_no: int):
        return None

    from src.chat.features.tutorial_search.services.document_chunking import (
        _pdf_to_markdown,
    )

    md = await _pdf_to_markdown(buf.getvalue(), vision_parser=failing_vision)
    assert "视觉解析不可用" in md


def test_parse_md_passthrough():
    assert asyncio.run(parse_to_markdown("a.md", "# 标题\n\n正文".encode())).startswith("# 标题")


def test_unsupported_extension_rejected():
    with pytest.raises(ValueError):
        asyncio.run(parse_to_markdown("virus.exe", b"binary"))
