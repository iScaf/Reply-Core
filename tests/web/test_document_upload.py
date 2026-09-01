# -*- coding: utf-8 -*-
"""教程上传端点端到端测试：multipart 上传 → 解析 → 切块 → 入库。

清理策略：快照精确删除——fixture 开始前记录已有文档 id 集合，
结束后只删除测试新创建的文档及其 chunks，绝不触碰库中已有数据
（tutorials 是知识库主数据，测试严禁整表清空）。
"""
import io

import pytest
from sqlalchemy import text

from tests.web.conftest import run_db


def _make_docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("上传部署手册", level=1)
    document.add_heading("第二章 环境准备", level=2)
    body = "这是一段很长的正文，用于确保章节超过切块阈值并拆分为父子结构。" * 36
    document.add_paragraph(body)
    document.add_heading("第三章 验证", level=2)
    document.add_paragraph("短章节内容。")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _snapshot_doc_ids() -> set:
    """记录测试前已存在的文档 id 集合。"""

    def _do(factory):
        async def inner(session):
            rows = await session.execute(
                text("SELECT id FROM tutorials.tutorial_documents")
            )
            return {r[0] for r in rows.fetchall()}

        async def flow():
            async with factory() as session:
                return await inner(session)

        return flow()

    return run_db(_do)


def _delete_created_docs(created_ids: set) -> None:
    """只删除指定 id 的文档及其 chunks（测试自己创建的）。"""
    if not created_ids:
        return
    id_list = ", ".join(str(int(i)) for i in created_ids)

    def _do(factory):
        async def inner(session):
            await session.execute(
                text(
                    f"DELETE FROM tutorials.knowledge_chunks "
                    f"WHERE document_id IN ({id_list})"
                )
            )
            await session.execute(
                text(
                    f"DELETE FROM tutorials.tutorial_documents "
                    f"WHERE id IN ({id_list})"
                )
            )

        async def flow():
            async with factory() as session:
                async with session.begin():
                    await inner(session)

        return flow()

    run_db(_do)


@pytest.fixture
def clean_tutorials_web():
    before = _snapshot_doc_ids()
    yield
    # 找出测试期间新创建的文档 id（含因 RESTART IDENTITY/自增产生的新增），
    # 只精确删除这些；已有数据一律不动
    after = _snapshot_doc_ids()
    _delete_created_docs(after - before)


@pytest.mark.usefixtures("clean_tutorials_web")
def test_upload_docx_end_to_end(web_client):
    resp = web_client.post(
        "/api/documents/tutorials/upload",
        files={"file": ("部署手册.docx", _make_docx_bytes(), "application/octet-stream")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["title"] == "部署手册"
    assert data["chunk_count"] >= 1
    assert data["doc_id"] > 0

    # 详情接口能看到切块结果
    detail = web_client.get(f"/api/documents/tutorials/{data['doc_id']}")
    assert detail.status_code == 200
    chunks = detail.json()["chunks"]
    assert len(chunks) == data["chunk_count"] + data["parent_count"]


@pytest.mark.usefixtures("clean_tutorials_web")
def test_upload_rejects_unsupported_format(web_client):
    resp = web_client.post(
        "/api/documents/tutorials/upload",
        files={"file": ("bad.exe", b"binary", "application/octet-stream")},
    )
    assert resp.status_code == 400
    assert "不支持" in resp.json()["detail"]


@pytest.mark.usefixtures("clean_tutorials_web")
def test_upload_rejects_oversize(web_client):
    big = b"x" * (10 * 1024 * 1024 + 1)
    resp = web_client.post(
        "/api/documents/tutorials/upload",
        files={"file": ("big.md", big, "text/markdown")},
    )
    assert resp.status_code == 400


@pytest.mark.usefixtures("clean_tutorials_web")
def test_update_reupload_delete_lifecycle(web_client):
    """编辑 → 重新上传 → 删除 的完整生命周期。"""
    # 准备：上传一篇
    resp = web_client.post(
        "/api/documents/tutorials/upload",
        files={"file": ("原始.md", ("# 原标题\n\n" + "原始正文内容。" * 100).encode(), "text/markdown")},
    )
    assert resp.status_code == 200
    doc_id = resp.json()["doc_id"]

    # 编辑：改标题与正文
    resp = web_client.put(
        f"/api/documents/tutorials/{doc_id}",
        json={"title": "新标题", "content": "# 新标题\n\n" + "编辑后的全新正文，长度足够触发切块。" * 60},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["title"] == "新标题"
    assert data["chunk_count"] >= 1

    # 详情确认编辑生效，且旧 chunk 被清掉（无残留"原始正文内容"）
    detail = web_client.get(f"/api/documents/tutorials/{doc_id}").json()
    assert detail["title"] == "新标题"
    assert all("原始正文内容" not in c["chunk_text"] for c in detail["chunks"])

    # 重新上传：换内容，标题保持
    resp = web_client.post(
        f"/api/documents/tutorials/{doc_id}/reupload",
        files={"file": ("替换版.md", ("# 替换后的内容标题\n\n" + "替换版正文。" * 80).encode(), "text/markdown")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["title"] == "新标题"  # 标题不变
    detail = web_client.get(f"/api/documents/tutorials/{doc_id}").json()
    assert "替换后的内容标题" in detail["full_text"]

    # 删除
    resp = web_client.delete(f"/api/documents/tutorials/{doc_id}")
    assert resp.status_code == 200
    assert web_client.get(f"/api/documents/tutorials/{doc_id}").status_code == 404


@pytest.mark.usefixtures("clean_tutorials_web")
def test_update_validates_empty_fields(web_client):
    resp = web_client.post(
        "/api/documents/tutorials/upload",
        files={"file": ("a.md", ("# 内容\n\n" + "这是一段足够长的正文内容。" * 30).encode(), "text/markdown")},
    )
    doc_id = resp.json()["doc_id"]

    resp = web_client.put(
        f"/api/documents/tutorials/{doc_id}", json={"title": "", "content": "x"}
    )
    assert resp.status_code == 400
    resp = web_client.put(
        f"/api/documents/tutorials/{doc_id}", json={"title": "t", "content": "  "}
    )
    assert resp.status_code == 400


@pytest.mark.usefixtures("clean_tutorials_web")
def test_delete_missing_returns_404(web_client):
    resp = web_client.delete("/api/documents/tutorials/999999")
    assert resp.status_code == 404
